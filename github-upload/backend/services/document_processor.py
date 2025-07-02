import os
import pandas as pd
from typing import Optional
import pytesseract
from PIL import Image
import PyPDF2
import pdfplumber
from docx import Document
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Processa diferentes tipos de documentos e extrai texto"""
    
    def __init__(self):
        # Configurar Tesseract se estiver no Windows
        if os.name == 'nt':
            # Caminho comum do Tesseract no Windows
            pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

    def process_document(self, file_path: str, content_type: str) -> str:
        """
        Processa documento baseado no tipo de arquivo
        """
        try:
            if content_type == "application/pdf":
                return self._process_pdf(file_path)
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._process_docx(file_path)
            elif content_type == "text/plain":
                return self._process_txt(file_path)
            elif content_type in ["text/csv", "application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
                return self._process_spreadsheet(file_path)
            elif content_type in ["image/png", "image/jpeg", "image/jpg"]:
                return self._process_image_ocr(file_path)
            else:
                raise ValueError(f"Tipo de arquivo não suportado: {content_type}")
                
        except Exception as e:
            logger.error(f"Erro ao processar documento {file_path}: {str(e)}")
            raise

    def _process_pdf(self, file_path: str) -> str:
        """Extrai texto de arquivo PDF"""
        text = ""
        try:
            # Tentar com pdfplumber primeiro (melhor para tabelas)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except:
            # Fallback para PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e:
                logger.error(f"Erro ao processar PDF: {str(e)}")
                
        return text.strip()

    def _process_docx(self, file_path: str) -> str:
        """Extrai texto de arquivo DOCX"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extrair texto de tabelas também
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
                    
            return text.strip()
        except Exception as e:
            logger.error(f"Erro ao processar DOCX: {str(e)}")
            return ""

    def _process_txt(self, file_path: str) -> str:
        """Lê arquivo de texto simples"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Tentar com encoding latin-1
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Erro ao processar TXT: {str(e)}")
                return ""

    def _process_spreadsheet(self, file_path: str) -> str:
        """Converte planilha em texto estruturado"""
        try:
            # Detectar tipo de arquivo
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Converter para texto estruturado
            text = f"Planilha com {len(df)} linhas e {len(df.columns)} colunas:\n\n"
            text += "Cabeçalhos: " + ", ".join(df.columns.tolist()) + "\n\n"
            
            # Adicionar algumas linhas de exemplo
            for idx, row in df.head(10).iterrows():
                text += f"Linha {idx + 1}: "
                for col in df.columns:
                    text += f"{col}: {row[col]}, "
                text = text.rstrip(", ") + "\n"
            
            if len(df) > 10:
                text += f"\n... e mais {len(df) - 10} linhas"
                
            return text
        except Exception as e:
            logger.error(f"Erro ao processar planilha: {str(e)}")
            return ""

    def _process_image_ocr(self, file_path: str) -> str:
        """Extrai texto de imagem usando OCR"""
        try:
            # Abrir imagem
            image = Image.open(file_path)
            
            # Usar Tesseract para OCR
            text = pytesseract.image_to_string(image, lang='por+eng')
            
            if not text.strip():
                logger.warning(f"Nenhum texto extraído da imagem {file_path}")
                return "Imagem processada, mas nenhum texto foi detectado."
            
            return text.strip()
        except Exception as e:
            logger.error(f"Erro ao processar imagem OCR: {str(e)}")
            return f"Erro ao processar imagem: {str(e)}"

    def get_document_metadata(self, file_path: str) -> dict:
        """Extrai metadados do documento"""
        metadata = {
            "filename": os.path.basename(file_path),
            "size": os.path.getsize(file_path),
            "extension": os.path.splitext(file_path)[1]
        }
        
        try:
            if file_path.endswith('.pdf'):
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata.update({
                        "pages": len(pdf_reader.pages),
                        "title": pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                        "author": pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else ''
                    })
        except:
            pass
            
        return metadata 